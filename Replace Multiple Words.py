from docx import Document
import openpyxl as px
import tkinter as tk
import tkinter.filedialog as fd
from tkinter import messagebox

#ウィンドウの準備
main_win = tk.Tk()
main_win.title("Excelリストに従って単語を置換する")
main_win.geometry("500x200")

main_frm = tk.Frame(main_win)
main_frm.grid(column=0, row=0, sticky=tk.NSEW, padx=5, pady=10)

def openList():
    list_path = fd.askopenfilename(filetypes=[('Excel','*.xlsx')])

    if list_path:
        list_filepath.set(list_path)
        #エクセルの読み込み
        wb = px.load_workbook(filename=list_path, data_only=True)
        ws = wb.active
        
        #検索する文字列リストの作成（空白行はエラー）
        source_words = []
        max=ws.max_row
        for cell in ws.iter_rows(max_row=max):
            for cell in ws['A']:
                if cell.value is None:
                    messagebox.showerror("Error", "空白のセルが検出されたため、リストを読み込めませんでした。\nリストを確認して修正してください。")
                    list_filepath.set("")
                    return
                source_words.append(cell.value)
        
        #置換後の文字列リストの作成（空白行はエラー）
        target_words = []
        max=ws.max_row
        for cell in ws.iter_rows(max_row=max):
            for cell in ws['B']:
                if cell.value is None:
                    messagebox.showerror("Error", "空白のセルが検出されたため、リストを読み込めませんでした。\nリストを確認して修正してください。")
                    list_filepath.set("")
                    return
                target_words.append(cell.value)

        #置換用ディクショナリの作成
        global dic
        dic = dict(zip(source_words,target_words))

    else:
        list_filepath.set("")
        messagebox.showinfo("Message", "単語リストを設定してください")

def openTarget():
    target_path = fd.askopenfilename(filetypes=[('Word','*.docx')])

    if target_path:
        target_filepath.set(target_path)
        #ワードの読み込み
        global dc
        dc = Document(target_path)

    else:
        target_filepath.set("")
        messagebox.showinfo("Message", "対象ファイルを設定してください")

def startRplace():
    #入力ボックスのチェック
    check_list = list_box.get()
    check_target = target_box.get()
    if check_list == "":
        messagebox.showerror("Error", "リストを設定してください")
    elif check_target == "":
        messagebox.showerror("Error", "対象ファイルを設定してください")
    else:
        output_file = fd.asksaveasfilename(
        filetypes=[("Word", "*.docx")], defaultextension=".docx"
        )
        count = 0

        if output_file:
            try:
                #置換件数をカウント
                for key,value in dic.items():
                    for par in dc.paragraphs:
                        word = par.text
                        count += word.count(key)

                if count > 0:
                    #ディクショナリを元に単語を置換
                    for key, value in dic.items():
                        for paragraph in dc.paragraphs:
                            if key in paragraph.text:
                                inline = paragraph.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, value)
                                        inline[i].text = text

                    #ワードを保存
                    dc.save(output_file)
                    messagebox.showinfo("Success", "完了しました！\n"+"置換件数："+str(count))
                else:
                    messagebox.showinfo("Message", "置換対象の単語はありませんでした")
            except:
                messagebox.showerror("Error", "エラーが発生したため、処理を中断しました。")
        else:
            messagebox.showerror("Error", "保存先とファイル名を設定してください")

list_filepath = tk.StringVar()
target_filepath = tk.StringVar()

#ウィジェット作成（リストパス）
list_label = tk.Label(main_frm, text="単語リスト")
list_box = tk.Entry(main_frm, textvariable=list_filepath)
list_btn = tk.Button(main_frm, text="参照", command = openList)

#ウィジェット作成（対象ファイルパス）
target_label = tk.Label(main_frm, text="対象ファイル")
target_box = tk.Entry(main_frm, textvariable=target_filepath)
target_btn = tk.Button(main_frm, text="参照", command = openTarget)

#ウィジェット作成（実行ボタン）
app_btn = tk.Button(main_frm, text="実行", command = startRplace)
app_btn.grid(column=1, row=3)

#入力ボックスのステータス
list_box.configure(state="readonly")
target_box.configure(state="readonly")

#配置
list_label.grid(column=0, row=0, pady=10)
list_box.grid(column=1, row=0, sticky=tk.EW, padx=5)
list_btn.grid(column=2, row=0)

target_label.grid(column=0, row=1)
target_box.grid(column=1, row=1, sticky=tk.EW, padx=5)
target_btn.grid(column=2, row=1)

main_win.columnconfigure(0, weight=1)
main_win.rowconfigure(0, weight=1)
main_frm.columnconfigure(1, weight=1)

main_win.mainloop()