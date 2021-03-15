from docx import Document
import openpyxl as px
import tkinter as tk
import tkinter.filedialog as fd
from tkinter import messagebox
import paragraph

#ウィンドウの準備
main_win = tk.Tk()
main_win.title("Excelリストに従って単語を置換する")
main_win.geometry("500x200")

main_frm = tk.Frame(main_win)
main_frm.grid(column=0, row=0, sticky=tk.NSEW, padx=5, pady=10)

def openList():
    list_path = fd.askopenfilename(filetypes=[('Excel','*.xlsx')])

    if list_path:
        list_filepath.set(list_path)
        #エクセルの読み込み
        wb = px.load_workbook(filename=list_path, data_only=True)
        ws = wb.active
        
        #検索する文字列リストの作成（空白行はエラー）
        source_words = []
        max=ws.max_row
        for cell in ws.iter_rows(max_row=max):
            for cell in ws['A']:
                if cell.value is None:
                    messagebox.showerror("Error", "空白のセルが検出されたため、リストを読み込めませんでした。\nリストを確認して修正してください。")
                    list_filepath.set("")
                    return
                source_words.append(cell.value)
        
        #置換後の文字列リストの作成（空白行はエラー）
        target_words = []
        max=ws.max_row
        for cell in ws.iter_rows(max_row=max):
            for cell in ws['B']:
                if cell.value is None:
                    messagebox.showerror("Error", "空白のセルが検出されたため、リストを読み込めませんでした。\nリストを確認して修正してください。")
                    list_filepath.set("")
                    return
                target_words.append(cell.value)

        #置換用ディクショナリの作成
        global dic
        dic = dict(zip(source_words,target_words))

    else:
        list_filepath.set("")
        messagebox.showinfo("Message", "単語リストを設定してください")

def openTarget():
    target_path = fd.askopenfilename(filetypes=[('Word','*.docx')])

    if target_path:
        target_filepath.set(target_path)
        #ワードの読み込み
        global dc
        dc = Document(target_path)

    else:
        target_filepath.set("")
        messagebox.showinfo("Message", "対象ファイルを設定してください")

def startRplace():
    #入力ボックスのチェック
    check_list = list_box.get()
    check_target = target_box.get()
    if check_list == "":
        messagebox.showerror("Error", "リストを設定してください")
    elif check_target == "":
        messagebox.showerror("Error", "対象ファイルを設定してください")
    else:
        output_file = fd.asksaveasfilename(
        filetypes=[("Word", "*.docx")], defaultextension=".docx"
        )
        count = 0

        if output_file:
            try:
                #置換件数をカウント
                for key,value in dic.items():
                    for para in dc.paragraphs:
                        word = para.text
                        count += word.count(key)

                if count > 0:
                    #ディクショナリを元に単語を置換
                    for key, value in dic.items():
                        for para in dc.paragraphs:
                            # ここから先、スペース区切りの用語も変更可能にするためにロジックを変更
                            found_pos = -1
                            key_count = para.text.count(key)
                            for _ in range(key_count):
                                # paragraphの中にキー（置換元）が何回出現するかを確認し、
                                # 出現回数分だけ置換処理を行う
                                #
                                # 以下、置換処理
                                # paragraph内でキーの開始位置とキーの終了位置を算出し、
                                # その位置情報を元にキーが含まれているrunを抽出する
                                # キーは複数のrunにまたがって存在する場合がある
                                # （スペース区切りの用語など）ためリストに抽出するようにする
                                found_pos = para.text.find(key, found_pos + 1)
                                if found_pos < 0: break
                                key_end_pos = found_pos + len(key) - 1
                                partial_para_text = ""
                                partial_runs = []
                                for run in para.runs:
                                    partial_para_text += run.text
                                    current_pos = len(partial_para_text) - 1
                                    if current_pos >= found_pos and current_pos < key_end_pos:
                                        partial_runs.append(run)
                                    elif current_pos >= key_end_pos:
                                        partial_runs.append(run)
                                        break
                                # 抽出した run のリストのテキストをすべてつなげる。
                                partial_runs_text = ""
                                for run in partial_runs:
                                    partial_runs_text = partial_runs_text + run.text

                                # 文章の途中にキーが含まれている場合は、キーだけではなく
                                # 前後の文も含まれている可能性があるため、
                                # つなげたテキストの中から、
                                # - キーより前の文
                                # - キー
                                # - キーより後の文
                                # を、それぞれの位置を割り出す。
                                key_start_pos = partial_runs_text.find(key)
                                after_key_start_pos = len(partial_runs_text) - (key_start_pos + len(key))
                                if len(partial_runs) < 2:
                                    # 抽出した run が 1つだけだった場合の置換処理
                                    #
                                    # キーより前の文 + 置換後のテキスト + キーより後の文
                                    # で結合したテキストをrun.textに設定する
                                    run = partial_runs[0]
                                    run.text = run.text[:key_start_pos] + value + run.text[len(run.text) - after_key_start_pos:]
                                else:
                                    # 抽出した run が 2つ以上だった場合の置換処理
                                    #
                                    # - 抽出した最後の run にキーの一部分が含まれている場合は、その部分を削除してrun.textに設定する
                                    run = partial_runs.pop(-1)
                                    run.text = run.text[len(run.text)-after_key_start_pos:]

                                    # - 抽出した最初の run に、『キーより前の文 + 置換後のテキスト』で結合したテキストをrun.textに設定する
                                    run = partial_runs.pop(0)
                                    run.text = run.text[:key_start_pos] + value

                                    # - 抽出した最後の run でも最初の run でもない run は、run.text を空文字にした上でリストから削除する
                                    for i in reversed(range(len(partial_runs))):
                                        para._p.remove(partial_runs[i]._r)
                    #ワードを保存
                    dc.save(output_file)
                    messagebox.showinfo("Success", "完了しました！\n"+"置換件数："+str(count))
                else:
                    messagebox.showinfo("Message", "置換対象の単語はありませんでした")
            except:
                messagebox.showerror("Error", "エラーが発生したため、処理を中断しました。")
        else:
            messagebox.showerror("Error", "保存先とファイル名を設定してください")

list_filepath = tk.StringVar()
target_filepath = tk.StringVar()

#ウィジェット作成（リストパス）
list_label = tk.Label(main_frm, text="単語リスト")
list_box = tk.Entry(main_frm, textvariable=list_filepath)
list_btn = tk.Button(main_frm, text="参照", command = openList)

#ウィジェット作成（対象ファイルパス）
target_label = tk.Label(main_frm, text="対象ファイル")
target_box = tk.Entry(main_frm, textvariable=target_filepath)
target_btn = tk.Button(main_frm, text="参照", command = openTarget)

#ウィジェット作成（実行ボタン）
app_btn = tk.Button(main_frm, text="実行", command = startRplace)
app_btn.grid(column=1, row=3)

#入力ボックスのステータス
list_box.configure(state="readonly")
target_box.configure(state="readonly")

#配置
list_label.grid(column=0, row=0, pady=10)
list_box.grid(column=1, row=0, sticky=tk.EW, padx=5)
list_btn.grid(column=2, row=0)

target_label.grid(column=0, row=1)
target_box.grid(column=1, row=1, sticky=tk.EW, padx=5)
target_btn.grid(column=2, row=1)

main_win.columnconfigure(0, weight=1)
main_win.rowconfigure(0, weight=1)
main_frm.columnconfigure(1, weight=1)

main_win.mainloop()